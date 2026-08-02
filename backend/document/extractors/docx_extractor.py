"""OpenXML DOCX Document Extractor (`DOCXExtractor`).

Extracts paragraphs and tables from `.docx` documents using `python-docx` or native XML parsing.
"""

import io
from typing import BinaryIO
import xml.etree.ElementTree as ET
import zipfile

from backend.document.schemas.errors import DocumentDomainException, DocumentErrorCode

from .base import BaseExtractor, ExtractedContent, ExtractorCapability


class DOCXExtractor(BaseExtractor):
    """Word document (`.docx`) extractor."""

    @property
    def capability(self) -> ExtractorCapability:
        return ExtractorCapability(
            name="DOCXExtractor",
            supported_mimes={
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            },
            supported_extensions={".docx"},
            priority=20,
            enabled=True,
        )

    async def extract(
        self, stream: BinaryIO, filename: str, mime_type: str
    ) -> ExtractedContent:
        try:
            current_pos = stream.tell()
            stream.seek(0)
            docx_bytes = stream.read()
            stream.seek(current_pos)

            text_paragraphs: list[str] = []
            metadata: dict[str, Any] = {}

            # Attempt python-docx first, otherwise native ZIP/XML extraction
            try:
                import docx

                doc = docx.Document(io.BytesIO(docx_bytes))
                for p in doc.paragraphs:
                    if p.text and p.text.strip():
                        text_paragraphs.append(p.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(
                            cell.text.strip() for cell in row.cells if cell.text.strip()
                        )
                        if row_text:
                            text_paragraphs.append(row_text)
            except ImportError:
                # Native OpenXML ZIP parsing
                try:
                    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
                        with z.open("word/document.xml") as xml_file:
                            tree = ET.parse(xml_file)
                            root = tree.getroot()
                            # OpenXML namespace
                            ns = {
                                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                            }
                            for node in root.findall(".//w:t", namespaces=ns):
                                if node.text and node.text.strip():
                                    text_paragraphs.append(node.text.strip())
                except Exception as zip_err:
                    raise DocumentDomainException(
                        code=DocumentErrorCode.EXTRACT_001,
                        message=f"Failed to parse DOCX OpenXML structure: {zip_err}",
                        detail={"filename": filename, "error": str(zip_err)},
                    ) from zip_err

            full_text = "\n\n".join(text_paragraphs).strip()
            words = full_text.split()
            word_count = len(words)
            page_count = max(1, (len(full_text) + 2999) // 3000)

            return ExtractedContent(
                text=full_text,
                word_count=word_count,
                page_count=page_count,
                metadata=metadata,
                language="en",
                needs_ocr=False,
            )
        except Exception as e:
            if isinstance(e, DocumentDomainException):
                raise
            raise DocumentDomainException(
                code=DocumentErrorCode.EXTRACT_001,
                message=f"DOCX extraction failed: {e}",
                detail={"filename": filename, "error": str(e)},
            ) from e
