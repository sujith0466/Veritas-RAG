import os
import json
import random
from faker import Faker
from datetime import datetime, timedelta

fake = Faker()
Faker.seed(1337)
random.seed(1337)

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

DEPARTMENTS = [
    'hr', 'engineering', 'it', 'security', 'product', 
    'finance', 'legal', 'customer_support', 'sales', 'marketing'
]

FORMATS = {'md': 0.30, 'docx': 0.15, 'pdf': 0.15, 'txt': 0.10, 'html': 0.05, 'json': 0.05, 'yaml': 0.05, 'csv': 0.05, 'xml': 0.05, 'py': 0.05}

def get_format():
    fmt = random.choices(list(FORMATS.keys()), weights=list(FORMATS.values()), k=1)[0]
    if fmt == 'docx' and not HAS_DOCX: return 'md'
    if fmt == 'pdf' and not HAS_REPORTLAB: return 'md'
    return fmt

def generate_hr_template():
    return f"""Employee Handbook - {fake.year()}
Effective Date: {fake.date_this_year()}

1. Introduction
Welcome to the company! We are thrilled to have you here. Your employee ID is {fake.unique.random_number(digits=5)}.

2. PTO Policy
All employees are entitled to 20 days of Paid Time Off annually.
Manager approval is required for consecutive leaves exceeding 5 days.

3. Code of Conduct
We maintain a safe, inclusive workspace. Harassment of any kind will not be tolerated.

Contact HR at hr@{fake.domain_name()} for more info."""

def generate_engineering_template():
    return f"""CI/CD Deployment Runbook
Service: {fake.word().capitalize()}-Service
Owner: {fake.name()}

1. Pre-requisites
- Ensure all tests pass in GitHub Actions.
- Target branch: main

2. Deployment Steps
```bash
git pull origin main
docker-compose build
docker-compose up -d
```

3. Rollback
If metrics drop below 99.9% uptime, run the rollback script immediately: `./scripts/rollback.sh {fake.hexify(text='^^^^')}`"""

def generate_it_template():
    return f"""VPN Setup Guide
Gateway: vpn.{fake.domain_name()}
Support Ticket Reference: IT-{fake.random_number(digits=4)}

Instructions:
1. Download the Cisco AnyConnect client.
2. Login with your corporate SSO credentials.
3. Approve the Duo Push notification on your mobile device.

Troubleshooting:
If you experience a "Connection Timeout", restart your router or contact IT Support at extension {fake.random_number(digits=3)}."""

def generate_security_template():
    return f"""Incident Response Protocol
Severity: Level 1 (Critical)
CISO: {fake.name()}

1. Detection
Monitor SIEM alerts for unauthorized access anomalies.

2. Containment
- Isolate the affected endpoints.
- Revoke compromised credentials immediately.

3. Eradication
Run malware scans and patch the exploited vulnerability (e.g., CVE-{fake.year()}-{fake.random_number(digits=4)}).

4. Post-Incident
Draft a retrospective report within 48 hours."""

def generate_finance_template():
    return f"""Expense Reimbursement Policy
Approved by: {fake.name()} (CFO)

1. Eligible Expenses
- Travel (economy class flights)
- Meals (up to $50 per diem)
- Client entertainment (requires pre-approval if >$200)

2. Submission
Submit all receipts via Concur within 30 days of the expense.

3. Approval Workflow
Employee -> Direct Manager -> Finance Team -> Payout."""

def generate_legal_template():
    return f"""Non-Disclosure Agreement (NDA)
Date: {fake.date_this_month()}
Between: Our Company and {fake.company()}

1. Confidential Information
"Confidential Information" means any proprietary data, financial records, or technical designs shared during the engagement.

2. Obligations
The Receiving Party shall not disclose the Confidential Information to any third party without prior written consent.

3. Term
This agreement shall remain in effect for {random.randint(1, 5)} years."""

def generate_generic_template(dept):
    return f"""{dept.capitalize()} Overview Document
Author: {fake.name()}
Date: {fake.date_this_year()}

1. Strategic Goals
Our primary objective for Q{random.randint(1,4)} is to increase efficiency by {random.randint(10,30)}%.

2. Key Performance Indicators (KPIs)
- Customer Satisfaction Score (CSAT): >90%
- Revenue Growth: {random.randint(5,15)}% YoY

3. Action Items
- Schedule weekly syncs.
- Review progress on {fake.date_this_month()}."""

TEMPLATES = {
    'hr': generate_hr_template,
    'engineering': generate_engineering_template,
    'it': generate_it_template,
    'security': generate_security_template,
    'finance': generate_finance_template,
    'legal': generate_legal_template,
}

out_dir = 'enterprise_demo_dataset'
os.makedirs(out_dir, exist_ok=True)
for d in DEPARTMENTS: os.makedirs(os.path.join(out_dir, d), exist_ok=True)

for i in range(120):
    dept = random.choice(DEPARTMENTS)
    fmt = get_format()
    
    # Generate content
    if dept in TEMPLATES:
        content_text = TEMPLATES[dept]()
    else:
        content_text = generate_generic_template(dept)
        
    lines = content_text.split('\n')
    title = lines[0] if lines else f"{dept.capitalize()} Document {i}"
    paragraphs = lines[1:]
    
    slug = title.lower().replace(' ', '_').replace('/', '_').replace('-', '_').replace(':', '')[:30] + '_' + str(i)
    path = os.path.join(out_dir, dept, f'{slug}.{fmt}')
    
    if fmt == 'pdf':
        doc = SimpleDocTemplate(path, pagesize=letter)
        story = [Paragraph(title, getSampleStyleSheet()['Heading1'])]
        for p in paragraphs:
            if p.strip():
                story.append(Paragraph(p, getSampleStyleSheet()['Normal']))
                story.append(Spacer(1, 10))
        doc.build(story)
    elif fmt == 'docx':
        doc = Document()
        doc.add_heading(title, 0)
        for p in paragraphs: 
            if p.strip(): doc.add_paragraph(p)
        doc.save(path)
    elif fmt == 'json':
        with open(path, 'w', encoding='utf-8') as f: json.dump({'title': title, 'content': paragraphs}, f)
    elif fmt == 'yaml':
        with open(path, 'w', encoding='utf-8') as f: f.write(f'title: {title}\ncontent:\n' + '\n'.join([f'  - {p}' for p in paragraphs if p.strip()]))
    elif fmt == 'csv':
        with open(path, 'w', encoding='utf-8') as f: f.write('ID,Content\n' + '\n'.join([f'{j},{p.replace(",", ";")}' for j, p in enumerate(paragraphs) if p.strip()]))
    elif fmt == 'xml':
        with open(path, 'w', encoding='utf-8') as f: f.write(f'<document><title>{title}</title><content>' + ''.join([f'<p>{p}</p>' for p in paragraphs if p.strip()]) + '</content></document>')
    elif fmt == 'html':
        with open(path, 'w', encoding='utf-8') as f: f.write(f'<html><head><title>{title}</title></head><body><h1>{title}</h1>' + ''.join([f'<p>{p}</p>' for p in paragraphs if p.strip()]) + '</body></html>')
    elif fmt == 'py':
        with open(path, 'w', encoding='utf-8') as f: f.write(f'\"\"\"{title}\"\"\"\n\ndef get_data():\n' + '\n'.join([f'    # {p}' for p in paragraphs if p.strip()]) + '\n    return True')
    else:
        with open(path, 'w', encoding='utf-8') as f: f.write(content_text)

print(f'Successfully generated 120 deterministic enterprise files in {out_dir}/')
