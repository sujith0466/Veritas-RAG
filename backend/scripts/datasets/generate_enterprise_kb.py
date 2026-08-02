import argparse
import asyncio
import csv
import io
import json
from pathlib import Path
import random
import time
import uuid

from faker import Faker

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import (
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from sqlalchemy import func, select
import structlog

from backend.database.engine import get_session_factory
from backend.document.models.document import Document
from backend.document.models.job import ProcessingJob
from backend.document.services.document_service import DocumentService
from backend.models.entities.user import User
from backend.modules.chunking.models.chunk import DocumentChunk
from backend.modules.embedding.models.chunk_embedding import ChunkEmbedding

logger = structlog.get_logger(__name__)
fake = Faker()

DEPARTMENTS = [
    "hr", "engineering", "it", "security", "product",
    "finance", "legal", "customer_support", "sales", "marketing"
]

# Weights for formats
FORMATS = {
    "md": 0.35,
    "docx": 0.20,
    "pdf": 0.15,
    "txt": 0.10,
    "html": 0.05,
    "json": 0.05,
    "yaml": 0.05,
    "csv": 0.03,
    "xml": 0.02
}

def get_weighted_format():
    formats = list(FORMATS.keys())
    weights = list(FORMATS.values())
    f = random.choices(formats, weights=weights, k=1)[0]
    if f == "pdf" and not HAS_REPORTLAB:
        return "md"
    if f == "docx" and not HAS_DOCX:
        return "md"
    return f

def generate_pdf(filepath, title, paragraphs):
    if not HAS_REPORTLAB:
        return
    doc = SimpleDocTemplate(str(filepath), pagesize=letter)
    styles = getSampleStyleSheet()
    title_style = styles['Heading1']
    normal_style = styles['Normal']

    story = []
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.2 * 72))

    for p in paragraphs:
        if p.startswith("# "):
            story.append(Paragraph(p[2:], styles['Heading2']))
        elif p.startswith("## "):
            story.append(Paragraph(p[3:], styles['Heading3']))
        else:
            story.append(Paragraph(p, normal_style))
        story.append(Spacer(1, 0.1 * 72))

    doc.build(story)

def generate_docx(filepath, title, paragraphs):
    if not HAS_DOCX:
        return
    doc = DocxDocument()
    doc.add_heading(title, 0)
    for p in paragraphs:
        if p.startswith("# "):
            doc.add_heading(p[2:], level=1)
        elif p.startswith("## "):
            doc.add_heading(p[3:], level=2)
        else:
            doc.add_paragraph(p)
    doc.save(str(filepath))

def create_document_content(department, size):
    title = f"{department.capitalize()} Document {fake.word().capitalize()} {fake.year()}"
    paragraphs = []

    # Generate metadata header
    meta = f"""---
title: {title}
department: {department}
author: {fake.name()}
date: {fake.date_this_decade()}
classification: {random.choice(['Internal', 'Confidential', 'Public'])}
---
"""
    num_paras = random.randint(3, 8) if size == "small" else (random.randint(10, 20) if size == "medium" else random.randint(30, 60))

    paras = [meta]
    for _ in range(num_paras):
        if random.random() < 0.2:
            paras.append(f"## {fake.catch_phrase()}")
        elif random.random() < 0.1 and department in ["engineering", "it"]:
            paras.append(f"```python\n# Auto-generated code\ndef {fake.word()}():\n    return '{fake.word()}'\n```")
        elif random.random() < 0.1:
            paras.append(f"- {fake.sentence()}\n- {fake.sentence()}\n- {fake.sentence()}")
        else:
            paras.append(fake.paragraph(nb_sentences=random.randint(3, 10)))

    # Add a cross-reference occasionally
    if random.random() < 0.3:
        ref_dep = random.choice(DEPARTMENTS)
        paras.append(f"For more information, please refer to the {ref_dep.capitalize()} policies.")

    return title, paras

def generate_file(out_dir, department, size, fmt):
    title, paragraphs = create_document_content(department, size)
    filename = f"{department}_{uuid.uuid4().hex[:8]}.{fmt}"
    filepath = out_dir / department / filename
    filepath.parent.mkdir(parents=True, exist_ok=True)

    text_content = "\n\n".join(paragraphs)
    mime_type = "text/plain"

    if fmt == "md":
        mime_type = "text/markdown"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(f"# {title}\n\n{text_content}")
    elif fmt == "txt":
        mime_type = "text/plain"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(f"{title}\n\n{text_content}")
    elif fmt == "html":
        mime_type = "text/html"
        html = f"<html><head><title>{title}</title></head><body><h1>{title}</h1>"
        for p in paragraphs:
            html += f"<p>{p}</p>"
        html += "</body></html>"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(html)
    elif fmt == "json":
        mime_type = "application/json"
        data = {
            "title": title,
            "department": department,
            "content": paragraphs
        }
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)
    elif fmt == "yaml":
        mime_type = "application/x-yaml"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(f"title: {title}\ndepartment: {department}\ncontent:\n")
            for p in paragraphs:
                f.write(f"  - {json.dumps(p)}\n")
    elif fmt == "csv":
        mime_type = "text/csv"
        with open(filepath, "w", encoding="utf-8", newline='') as f:
            writer = csv.writer(f)
            writer.writerow(["Title", "Content"])
            for p in paragraphs:
                writer.writerow([title, p])
    elif fmt == "xml":
        mime_type = "application/xml"
        xml = f"<?xml version=\"1.0\" encoding=\"UTF-8\"?>\n<document><title>{title}</title><department>{department}</department><content>"
        for p in paragraphs:
            xml += f"<paragraph>{p}</paragraph>"
        xml += "</content></document>"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(xml)
    elif fmt == "pdf":
        mime_type = "application/pdf"
        generate_pdf(filepath, title, paragraphs)
    elif fmt == "docx":
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        generate_docx(filepath, title, paragraphs)

    return filepath, filename, mime_type

async def async_main(workspace_email: str, count: int):
    logger.info("Initializing Enterprise KB Generation", workspace=workspace_email, target_count=count)

    session_maker = get_session_factory()
    async with session_maker() as session:
        # Resolve user
        stmt = select(User).where(User.email == workspace_email)
        result = await session.execute(stmt)
        user = result.scalar_one_or_none()

        if not user:
            logger.error("Workspace user not found.", email=workspace_email)
            return

        tenant_id = user.tenant_id
        owner_id = user.id

        logger.info("Resolved tenant", tenant_id=tenant_id, user_id=owner_id)

        out_dir = Path("backend/demo_data")
        out_dir.mkdir(parents=True, exist_ok=True)

        doc_service = DocumentService()

        # Seed for idempotency
        random.seed(42)
        Faker.seed(42)
        global fake
        fake = Faker()

        # Check existing docs for idempotency
        existing_stmt = select(Document.filename).where(Document.tenant_id == tenant_id)
        existing_files = {row[0] for row in (await session.execute(existing_stmt)).all()}
        logger.info("Found existing documents in tenant", count=len(existing_files))
        logger.info("Found existing documents in tenant", count=len(existing_files))

        generated_count = 0
        skipped_count = 0

        # Generation phase
        files_to_upload = []

        for i in range(count):
            dept = random.choice(DEPARTMENTS)
            fmt = get_weighted_format()
            size = random.choices(["small", "medium", "large"], weights=[0.5, 0.3, 0.2])[0]

            filepath, filename, mime_type = generate_file(out_dir, dept, size, fmt)

            if filename in existing_files:
                skipped_count += 1
                continue

            files_to_upload.append((filepath, filename, mime_type))
            generated_count += 1
            if generated_count % 50 == 0:
                logger.info("Generated batch", count=generated_count)

        logger.info("Generation complete", generated=generated_count, skipped=skipped_count)

        # Upload phase
        logger.info("Starting uploads to pipeline...")
        for filepath, filename, mime_type in files_to_upload:
            try:
                with open(filepath, "rb") as f:
                    stream = io.BytesIO(f.read())

                await doc_service.upload_document(
                    stream=stream,
                    filename=filename,
                    declared_mime=mime_type,
                    tenant_id=tenant_id,
                    owner_user_id=owner_id,
                    session=session
                )
                await session.commit()
            except Exception as e:
                logger.error("Failed to upload document", filename=filename, error=str(e))
                await session.rollback()

        logger.info("Uploads complete. Waiting for processing to finish (this may take a few minutes)...")

        # Polling for completion
        start_time = time.time()
        timeout = 600  # 10 minutes
        while time.time() - start_time < timeout:
            pending_stmt = select(func.count(ProcessingJob.id)).join(
                Document, ProcessingJob.document_id == Document.id
            ).where(
                Document.tenant_id == tenant_id,
                ProcessingJob.status.in_(["PENDING", "EXTRACTING", "VALIDATING", "NORMALIZING", "CHUNKING", "EMBEDDING"])
            )
            pending_count = (await session.execute(pending_stmt)).scalar() or 0

            if pending_count == 0:
                break

            logger.info("Processing pipeline active...", pending_jobs=pending_count)
            await asyncio.sleep(10)

        elapsed = time.time() - start_time
        logger.info("Pipeline processing finished or timed out.", elapsed_seconds=elapsed)

        # Statistics Gathering
        print("\n==================================================")
        print("ENTERPRISE KNOWLEDGE BASE SEEDING REPORT")
        print("==================================================")

        # Docs
        res = await session.execute(select(Document.status, func.count(Document.id)).where(Document.tenant_id == tenant_id).group_by(Document.status))
        doc_stats = dict(res.all())
        total_docs = sum(doc_stats.values())
        print(f"Total Documents: {total_docs}")
        for k, v in doc_stats.items():
            print(f"  - {k}: {v}")

        # Chunks
        res = await session.execute(select(func.count(DocumentChunk.id)).where(DocumentChunk.tenant_id == tenant_id))
        total_chunks = res.scalar() or 0
        print(f"\nChunks Created: {total_chunks}")

        # Embeddings
        res = await session.execute(select(func.count(ChunkEmbedding.id)).where(ChunkEmbedding.tenant_id == tenant_id))
        total_embeddings = res.scalar() or 0
        print(f"Embeddings Created: {total_embeddings}")

        # Vectors - Note: Vector sync runs on a schedule or task, we can approximate by embeddings
        # Assuming Vectors Stored = Embeddings Created
        print(f"Vectors Stored: {total_embeddings}")

        # Jobs
        res = await session.execute(
            select(ProcessingJob.status, func.count(ProcessingJob.id))
            .join(Document, ProcessingJob.document_id == Document.id)
            .where(Document.tenant_id == tenant_id)
            .group_by(ProcessingJob.status)
        )
        job_stats = dict(res.all())
        print("\nJobs Status:")
        for k, v in job_stats.items():
            print(f"  - {k}: {v}")

        print(f"\nTotal Processing Time: {elapsed:.2f} seconds")

        # Top 10 documents by size
        res = await session.execute(
            select(Document.filename, Document.file_size_bytes)
            .where(Document.tenant_id == tenant_id)
            .order_by(Document.file_size_bytes.desc())
            .limit(10)
        )
        print("\nTop 10 Document Sizes:")
        for doc_name, size in res.all():
            print(f"  - {doc_name}: {size} bytes")

        # Top 10 generated documents (latest 10)
        res = await session.execute(
            select(Document.filename)
            .where(Document.tenant_id == tenant_id)
            .order_by(Document.created_at.desc())
            .limit(10)
        )
        print("\nTop 10 Generated Document Names:")
        for (doc_name,) in res.all():
            print(f"  - {doc_name}")

        print("==================================================\n")

def main():
    parser = argparse.ArgumentParser(description="Generate Enterprise Knowledge Base for RAGuard")
    parser.add_argument("--workspace", type=str, required=True, help="Email of the workspace admin to seed")
    parser.add_argument("--documents", type=int, default=200, help="Number of documents to generate")
    args = parser.parse_args()

    asyncio.run(async_main(args.workspace, args.documents))

if __name__ == "__main__":
    main()
